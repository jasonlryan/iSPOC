#!/usr/bin/env python3
"""
Script to convert DOCX guide files to structured JSON format
"""

import os
import json
import re
import docx
from datetime import datetime

# Paths
INPUT_DIR = "raw_guides"
OUTPUT_DIR = "VECTOR_GUIDES_JSON"

# Ensure output directory exists
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Helper function to extract guide information from filename
def extract_guide_info(filename):
    # Extract guide number and title
    match = re.match(r'(\d+)\.?\s+(.+)\.docx', filename)
    if match:
        guide_number, title = match.groups()
        return guide_number, title
    return None, filename.replace('.docx', '')

# Helper function to clean and structure text
def clean_text(text):
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# Function to find section headers in document
def identify_sections(paragraphs):
    sections = {
        "overview": "",
        "steps": "",
        "prerequisites": "",
        "troubleshooting": "",
        "examples": "",
        "notes": ""
    }
    
    current_section = "overview"  # Default section
    section_text = []
    
    # Common section header patterns for guides
    section_patterns = {
        r'(?i)overview|introduction|summary': "overview",
        r'(?i)steps|procedure|instruction|how to': "steps",
        r'(?i)prerequisite|before you begin|requirements': "prerequisites",
        r'(?i)troubleshoot|problems|issues|errors': "troubleshooting",
        r'(?i)example|sample': "examples",
        r'(?i)notes|tip|additional information': "notes"
    }
    
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this paragraph is a section header
        is_section_header = False
        for pattern, section_name in section_patterns.items():
            if re.search(pattern, text) and (para.style.name.startswith('Heading') or all(run.bold for run in para.runs)):
                current_section = section_name
                section_text = []
                is_section_header = True
                break
                
        if not is_section_header:
            section_text.append(text)
            sections[current_section] = " ".join(section_text)
    
    return {k: clean_text(v) for k, v in sections.items() if v}

# Function to process a single document
def process_document(file_path):
    try:
        # Get filename without path
        filename = os.path.basename(file_path)
        
        # Skip non-docx files
        if not filename.endswith('.docx'):
            return None
            
        # Extract guide number and title
        guide_number, title = extract_guide_info(filename)
        
        # Parse document
        doc = docx.Document(file_path)
        
        # Get document properties
        doc_properties = {}
        for prop in doc.core_properties.__dict__.items():
            if prop[0].startswith('_'):
                continue
            if prop[1] and hasattr(prop[1], 'strftime'):
                doc_properties[prop[0]] = prop[1].strftime('%Y-%m-%d')
            elif prop[1]:
                doc_properties[prop[0]] = str(prop[1])
        
        # Extract document text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Identify sections in the document
        sections = identify_sections(doc.paragraphs)
        
        # Create structured JSON
        guide_json = {
            "guide_number": guide_number if guide_number else "unknown",
            "title": title,
            "filename": filename,
            "extracted_date": datetime.now().strftime('%Y-%m-%d'),
            "metadata": doc_properties,
            "full_text": full_text,
            "sections": sections
        }
        
        # Create output filename
        output_file = os.path.join(OUTPUT_DIR, f"{filename.replace('.docx', '.json')}")
        
        # Write to JSON file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(guide_json, f, indent=2, ensure_ascii=False)
            
        return output_file
            
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        return None

# Main function
def main():
    print(f"Starting conversion of guide DOCX files to JSON format...")
    print(f"Input directory: {INPUT_DIR}")
    print(f"Output directory: {OUTPUT_DIR}")
    
    # Get list of all docx files
    docx_files = []
    for file in os.listdir(INPUT_DIR):
        file_path = os.path.join(INPUT_DIR, file)
        if file.endswith('.docx') and os.path.isfile(file_path):
            docx_files.append(file_path)
    
    print(f"Found {len(docx_files)} DOCX files to process.")
    
    # Process each file
    processed_files = 0
    for file_path in docx_files:
        print(f"Processing: {os.path.basename(file_path)}")
        output_file = process_document(file_path)
        if output_file:
            processed_files += 1
            print(f"  ✓ Created: {output_file}")
        else:
            print(f"  ✗ Failed to process")
    
    print(f"\nConversion complete. Processed {processed_files} of {len(docx_files)} files.")

if __name__ == "__main__":
    main() 