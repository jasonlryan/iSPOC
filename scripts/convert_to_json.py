#!/usr/bin/env python3
"""
Script to convert DOCX policy files to structured JSON format
"""

import os
import json
import re
import docx
from datetime import datetime

# Paths
INPUT_DIR = "raw policies"
OUTPUT_DIR = "VECTOR_JSON"

# Ensure output directory exists
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Helper function to extract policy information from filename
def extract_policy_info(filename):
    # Extract policy ID and title
    match = re.match(r'([A-Z]+\d+(?:\.\d+)?[a-z]*)(?:\s+)(.+)\.docx', filename)
    if match:
        policy_id, title = match.groups()
        return policy_id, title
    return None, filename.replace('.docx', '')

# Helper function to clean and structure text
def clean_text(text):
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# Function to find section headers in document
def identify_sections(paragraphs):
    sections = {
        "summary": "",
        "purpose": "",
        "scope": "",
        "definitions": "",
        "policy": "",
        "procedure": "",
        "responsibilities": "",
        "references": ""
    }
    
    current_section = "summary"  # Default section
    section_text = []
    
    # Common section header patterns
    section_patterns = {
        r'(?i)purpose|objective|aims': "purpose",
        r'(?i)scope|applies to|application': "scope",
        r'(?i)definition|terminology|terms used': "definitions",
        r'(?i)policy statement|policy|principle': "policy",
        r'(?i)procedure|process|method': "procedure",
        r'(?i)responsibilit|duties|roles': "responsibilities",
        r'(?i)reference|related document|further reading': "references",
        r'(?i)introduction|summary|overview': "summary"
    }
    
    for para in paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this paragraph is a section header
        is_section_header = False
        for pattern, section_name in section_patterns.items():
            if re.search(pattern, text) and (para.style.name.startswith('Heading') or all(run.bold for run in para.runs)):
                current_section = section_name
                section_text = []
                is_section_header = True
                break
                
        if not is_section_header:
            section_text.append(text)
            sections[current_section] = " ".join(section_text)
    
    return {k: clean_text(v) for k, v in sections.items() if v}

# Function to process a single document
def process_document(file_path):
    try:
        # Get filename without path
        filename = os.path.basename(file_path)
        
        # Skip non-docx files
        if not filename.endswith('.docx'):
            return None
            
        # Extract policy ID and title
        policy_id, title = extract_policy_info(filename)
        
        # Parse document
        doc = docx.Document(file_path)
        
        # Get document properties
        doc_properties = {}
        for prop in doc.core_properties.__dict__.items():
            if prop[0].startswith('_'):
                continue
            if prop[1] and hasattr(prop[1], 'strftime'):
                doc_properties[prop[0]] = prop[1].strftime('%Y-%m-%d')
            elif prop[1]:
                doc_properties[prop[0]] = str(prop[1])
        
        # Extract document text
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Identify sections in the document
        sections = identify_sections(doc.paragraphs)
        
        # Create structured JSON
        policy_json = {
            "id": policy_id if policy_id else "unknown",
            "title": title,
            "filename": filename,
            "extracted_date": datetime.now().strftime('%Y-%m-%d'),
            "metadata": doc_properties,
            "full_text": full_text,
            "sections": sections
        }
        
        # Create output filename
        output_file = os.path.join(OUTPUT_DIR, f"{filename.replace('.docx', '.json')}")
        
        # Write to JSON file
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(policy_json, f, indent=2, ensure_ascii=False)
            
        return output_file
            
    except Exception as e:
        print(f"Error processing {file_path}: {str(e)}")
        return None

# Main function
def main():
    print(f"Starting conversion of DOCX files to JSON format...")
    print(f"Input directory: {INPUT_DIR}")
    print(f"Output directory: {OUTPUT_DIR}")
    
    # Get list of all docx files
    docx_files = []
    for file in os.listdir(INPUT_DIR):
        file_path = os.path.join(INPUT_DIR, file)
        if file.endswith('.docx') and os.path.isfile(file_path):
            docx_files.append(file_path)
    
    print(f"Found {len(docx_files)} DOCX files to process.")
    
    # Process each file
    processed_files = 0
    for file_path in docx_files:
        print(f"Processing: {os.path.basename(file_path)}")
        output_file = process_document(file_path)
        if output_file:
            processed_files += 1
            print(f"  ✓ Created: {output_file}")
        else:
            print(f"  ✗ Failed to process")
    
    print(f"\nConversion complete. Processed {processed_files} of {len(docx_files)} files.")

if __name__ == "__main__":
    main() 